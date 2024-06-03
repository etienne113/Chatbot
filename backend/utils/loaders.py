from docx import Document as DocxDocument


class Document:
    def __init__(self, page_content, metadata):
        self.page_content = page_content
        self.metadata = metadata


class DocxLoader:
    def __init__(self, file_path):
        self.file_path = file_path

    def load(self):
        document = DocxDocument(self.file_path)
        text = '\n'.join([para.text for para in document.paragraphs])
        return [Document(page_content=text, metadata={})]
